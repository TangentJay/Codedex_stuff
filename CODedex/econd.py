from docx import Document
from docx.shared import Inches
from datetime import datetime

# Create a new Word document
doc = Document()
doc.add_heading("The Four Factors of Production – AP Economics Cheatsheet", level=1)

# Add an introductory paragraph
doc.add_paragraph(
    "This cheatsheet covers the four factors of production with AP-level depth, "
    "including definitions, examples, economic returns, and comparisons for effective study."
)

# Add detailed table of the Four Factors of Production
doc.add_heading("Detailed Breakdown", level=2)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Factor of Production'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'Examples'
hdr_cells[3].text = 'Economic Return'
hdr_cells[4].text = 'Key Characteristics'

rows = [
    ["Land", "All natural resources used in production.", 
     "Soil, water, oil, minerals, timber, farmland.", "Rent",
     "• Naturally occurring\n• Limited supply\n• Renewable or not"],

    ["Labor", "Human physical or mental effort in production.", 
     "Teacher, worker, engineer, barista.", "Wages",
     "• Can improve via education\n• Varies in skill\n• Subject to productivity"],

    ["Capital", "Man-made tools and equipment used in production.", 
     "Machines, factories, software, computers.", "Interest",
     "• Increases efficiency\n• Must be produced\n• Depreciates over time"],

    ["Entrepreneurship", "The skill to organize the other factors and take risks.", 
     "Startup founder, app developer, bakery owner.", "Profit",
     "• Combines factors\n• Takes risks\n• Drives innovation"]
]

for row_data in rows:
    row_cells = table.add_row().cells
    for i, item in enumerate(row_data):
        row_cells[i].text = item

# Add Comparison Table
doc.add_heading("Quick Comparison Table", level=2)
table2 = doc.add_table(rows=1, cols=5)
hdr2 = table2.rows[0].cells
hdr2[0].text = "Aspect"
hdr2[1].text = "Land"
hdr2[2].text = "Labor"
hdr2[3].text = "Capital"
hdr2[4].text = "Entrepreneurship"

comparison_rows = [
    ["Resource Type", "Natural", "Human", "Man-made", "Human (organizational)"],
    ["Earns Return As", "Rent", "Wages", "Interest", "Profit"],
    ["Renewable?", "Sometimes", "Yes", "Yes", "Yes"],
    ["Example Input", "Oil, rivers", "Workers", "Equipment", "Business owner"],
    ["Depreciates?", "No", "No", "Yes", "No (but can fail)"],
    ["Key Role", "Base resource", "Labor power", "Efficient production", "Creates and manages"]
]

for row_data in comparison_rows:
    row_cells = table2.add_row().cells
    for i, item in enumerate(row_data):
        row_cells[i].text = item

# Add memory tips
doc.add_heading("Memory Tips (Mnemonics)", level=2)
doc.add_paragraph(
    "- Land → Location, Lakes, Lumber (Nature!)\n"
    "- Labor → Lifting, Learning, Lives (People!)\n"
    "- Capital → Cranes, Computers, Construction (Tools!)\n"
    "- Entrepreneurship → Energy, Effort, Execution (Vision!)"
)

# Add application example
doc.add_heading("Example Scenario", level=2)
doc.add_paragraph(
    "In a coffee shop:\n"
    "- Land: The lot where the shop is built\n"
    "- Labor: Baristas and cashiers\n"
    "- Capital: Espresso machines, furniture\n"
    "- Entrepreneur: Owner who organizes it all and takes the risk"
)

# Save the document
file_path = "/mnt/data/AP_Economics_Factors_of_Production_Cheatsheet.docx"
doc.save(file_path)

print(f"Document saved as {file_path}")
